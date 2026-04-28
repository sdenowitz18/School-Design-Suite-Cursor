#!/usr/bin/env python3
"""Build AI_Testing_Question_Bank.docx from AI_Testing_Question_Bank.md with Word formatting."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Pt

# Rotate through these so each new subsection restarts at 1 in Word.
LIST_NUMBER_STYLES = ("List Number", "List Number 2", "List Number 3")


def add_inline_runs(paragraph, text: str) -> None:
    """Split on **bold** and append runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part:
            paragraph.add_run(part)


def compact_list_paragraph(paragraph) -> None:
    """Tighter vertical spacing like Word's compact multilevel lists."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def add_word_numbered_item(doc: Document, list_style: str, body: str) -> None:
    """One list item; Word supplies the digit + indent (native 'indented numbers')."""
    p = doc.add_paragraph(style=list_style)
    compact_list_paragraph(p)
    add_inline_runs(p, body)


def add_manual_dot_numbered(doc: Document, num: str, body: str) -> None:
    """Hanging-indent paragraph for N. questions (sections broken by bold headers)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.36)
    pf.first_line_indent = Inches(-0.36)
    r1 = p.add_run(f"{num}.\u00a0")
    r1.bold = True
    add_inline_runs(p, body)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "AI_Testing_Question_Bank.md"
    out_path = root / "AI_Testing_Question_Bank.docx"

    if not md_path.is_file():
        print(f"Missing {md_path}", file=sys.stderr)
        return 1

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Subsections that use Word list numbering (restart when style rotates).
    list_subsection_index = -1
    active_list_style = LIST_NUMBER_STYLES[0]

    def bump_list_subsection() -> None:
        nonlocal list_subsection_index, active_list_style
        list_subsection_index += 1
        active_list_style = LIST_NUMBER_STYLES[list_subsection_index % len(LIST_NUMBER_STYLES)]

    lines = md_path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s == "---":
            continue

        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=3)
            continue

        if s.startswith("<u>") and s.endswith("</u>"):
            inner = s[3:-4]
            bump_list_subsection()
            p = doc.add_paragraph()
            compact_list_paragraph(p)
            r = p.add_run(inner)
            r.underline = WD_UNDERLINE.SINGLE
            continue

        if s.startswith("**") and s.endswith("**") and s.count("**") == 2:
            inner = s[2:-2]
            # New list stream for Ring snapshot (follows underlined center subsections).
            if inner.startswith("Ring Component"):
                bump_list_subsection()
            p = doc.add_paragraph()
            compact_list_paragraph(p)
            r = p.add_run(inner)
            r.bold = True
            continue

        if len(s) >= 2 and s.startswith("*") and s.endswith("*") and not s.startswith("* "):
            inner = s[1:-1]
            p = doc.add_paragraph()
            compact_list_paragraph(p)
            r = p.add_run(inner)
            r.italic = True
            continue

        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            compact_list_paragraph(p)
            add_inline_runs(p, s[2:])
            continue

        m_paren = re.match(r"^(\d+)\)\s+(.*)$", s)
        if m_paren:
            add_word_numbered_item(doc, active_list_style, m_paren.group(2))
            continue

        m_dot = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m_dot:
            add_manual_dot_numbered(doc, m_dot.group(1), m_dot.group(2))
            continue

        p = doc.add_paragraph()
        compact_list_paragraph(p)
        add_inline_runs(p, s)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
